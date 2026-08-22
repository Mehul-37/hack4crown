import io
from typing import List, Dict, Any
import docx

def process_docx(file_bytes: bytes, filename: str) -> List[Dict[str, Any]]:
    """
    Extract text, headings, paragraphs, and table contents from .docx file.
    Returns page structure [{"page": 1, "text": "..."}]
    """
    try:
        doc = docx.Document(io.BytesIO(file_bytes))
        extracted_lines = []
        
        # Extract paragraph text
        for para in doc.paragraphs:
            if para.text.strip():
                extracted_lines.append(para.text.strip())
                
        # Extract table content
        for table in doc.tables:
            for row in table.rows:
                row_cells = [cell.text.strip() for cell in row.cells if cell.text.strip()]
                if row_cells:
                    extracted_lines.append(" | ".join(row_cells))
                    
        full_text = "\n".join(extracted_lines).strip()
        if not full_text:
            full_text = f"Docx document {filename} contained no text content."
            
        return [{
            "page": 1,
            "text": full_text
        }]
    except Exception as e:
        return [{
            "page": 1,
            "text": f"Error parsing docx document {filename}: {str(e)}"
        }]
